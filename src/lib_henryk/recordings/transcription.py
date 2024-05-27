import requests
import pandas as pd
import os
import shutil
import io
import json
import math
import docx
import progress

import subprocess

# local imports
from lib_henryk.logger import *
from lib_henryk.config import *
from lib_henryk.utils import *


def submit_transcriptions_goodtape( df: pd.DataFrame, webhooks_token_id, goodtape_api_key, path_recordings, start=0, stop=None, verbose=False ):        
    # create new transcription_id column if needed
    if not 'transcription_id' in  df.columns.to_list():
        df['transcription_id'] = None

    # create new transcription_json column if needed
    if not 'transcription_json' in  df.columns.to_list():
        df['transcription_json'] = None

    # create new transcription_txt column if needed
    if not 'transcription_txt' in  df.columns.to_list():
        df['transcription_txt'] = None
    
    # set up callback
    token_id = webhooks_token_id
    callback_url = f'https://webhook.site/{token_id}'

    # correct stop and start
    start = start if (start is not None) and (start >= 0) and (start < len(df)) else 0
    stop = stop if (stop is not None) and (stop > 0) and (stop <= len(df)) and (stop <= start) else len(df)
    
    # load just few records
    counter=0
    for index, row in df[start:stop].iterrows():
        file_path = path_recordings + '/' + row['path']
        file_name = row['file']
        transcription_id = row['transcription_id']
        file_name_json = row['transcription_json']
        file_name_txt= row['transcription_txt']
    
        # run transcription API if needed, only if transcription wasn't done already
        # we are only checking for json file, because transcription_id won't be needed anymore
        if pd.isna(file_name_json) == True:
            file_name_max_length = 100
            file_name_truncated = file_name
            if len(file_name) > file_name_max_length:
                file_name_truncated = file_name[0:file_name_max_length] + "..."            
            print(f'processing [{index}]: {file_name_truncated}')

            # copy file to /tmp, curl has issue with some paths
            temp_ext = file_name.split('.')[-1]
            temp_path = f'/tmp/recording.{temp_ext}'
            shutil.copy(file_path, temp_path)

            # perform transcription via API & CURL        
            response = subprocess.check_output(
                f'curl -s -X POST "https://api.goodtape.io/transcribe"' \
                + f' -H "Authorization: {goodtape_api_key}"' \
                + f' -F "audio=@{temp_path}"' \
                + f' -F "callbackUrl={callback_url}"' \
                + f' -F "languageCode=pl"' \
                + f' -F "speakerLabels=true"' \
                + f' -F "timeStamps=false"' 
            , shell=True, ) 

            # verify and write down the response
            try:
                response_json = json.loads(response)
                transcription_id = response_json['transcriptionId']
                df.loc[index, 'transcription_id'] = transcription_id
                if verbose:
                    print(f'file [{index}]: {transcription_id}')
                counter+=1
            except:
                print(response)
                break
        else:
            if verbose:
                print(f'file [{index}]: {file_name} already submitted')
    
    print(f'submitted {counter} recordings to the transcription service GoodTape')


def retrieve_responses_goodtape_via_webhooks(df, webhooks_token_id, path_transcriptions_json):
    # create new column if needed
    if not 'transcription_json' in  df.columns.to_list():
        df['transcription_json'] = None
        
    # use webhooks.site API to pull latest requests list
    headers = {}
    r = requests.get('https://webhook.site/token/'+ webhooks_token_id +'/requests?sorting=newest', headers=headers)
    print(f'found {len(r.json()["data"])} requests to fetch')
    
    # process requests one by one
    counter=0
    for i, request in enumerate(r.json()['data']):
        # fetch the data         
        response = requests.get(f'https://webhook.site/token/{webhooks_token_id}/request/{request["uuid"]}/raw')
        json_content = response.json()
        transcription_id = json_content['transcription_id']

        # before we write the file, let's just check if
        # we have a corresponding transcription_id in the database
        idx = df[df['transcription_id']==transcription_id].index
        if idx == None: 
            continue
        else:
            file_name = df.loc[idx[0]]['file']
            file_index = idx[0]
            file_name_max_length = 64
            file_name_name_truncated = file_name
            file_name_without_ext = ".".join(file_name.split(".")[:-1])
            file_name_json = f'{file_name_without_ext}.json'
            path_json = f'{path_transcriptions_json}/{file_name_json}'
            if len(file_name) > file_name_max_length:
                file_name_truncated = file_name[0:file_name_max_length] + "..."
    
        # write json file
        with open(f'{path_json}', "w") as file: 
            json.dump(json_content, file, indent = 4) 
            file.close() 
    
        # mark as done in the transcription log
        df.loc[file_index, 'transcription_json'] = file_name_json
        print(f'retrieved transcription [{i}]: {transcription_id} -> file [{file_index}]: {file_name_truncated}')
        counter+=1
        
        # clean up this request as it was fetched
        response = requests.delete(f'https://webhook.site/token/{webhooks_token_id}/request/{request["uuid"]}')
    
    # print message that we are done
    print(f'{counter} transcriptions were retrieved')


def process_transcriptions_json(
    df, 
    path_transcriptions_json, 
    path_transcriptions_txt, 
    path_transcriptions_doc, 
    to_txt=True, 
    to_doc=True, 
    verbose=False
):
    # create new column if needed
    if (not 'transcription_txt' in  df.columns.to_list()):
        df['transcription_txt'] = None
    if (not 'transcription_doc' in  df.columns.to_list()):
        df['transcription_doc'] = None
    
    # iterate over records and process files
    counter=0
    for index, row in df.iterrows():
        progressBar(index, len(df), prefix=f'processing {len(df)} transcriptions')
        
        file_name = row['file']
        file_name_without_ext = get_file_name_without_extension(file_name)
        file_name_json = row['transcription_json']
        path_json = f'{path_transcriptions_json}/{file_name_json}'

        # doc and txt file names and paths
        file_name_txt = f'{file_name_without_ext}.txt'
        file_name_doc = f'{file_name_without_ext}.docx'
        path_txt = f'{path_transcriptions_txt}/{file_name_txt}'
        path_doc = f'{path_transcriptions_doc}/{file_name_doc}'

        # failover in case json is not available
        if pd.isna(file_name_json) or (os.path.exists(path_json) == False): 
            continue
    
        # open json file
        transcription_text = ''
        with open(path_json, 'r') as file:
            json_content = json.load(file)
            transcription_text = json_content['content']['text']
            file.close()
        
        # if to_txt hasn't been done already
        processed = False
        if to_txt and (not os.path.exists(path_txt)):
            with open(path_txt, 'w') as file:
                file.write(file_name_without_ext + '\n\n')
                file.write(transcription_text.strip())
                file.close()    
            df.loc[index, 'transcription_txt'] = file_name_txt
            processed = True
        elif to_txt and os.path.exists(path_txt):
            if verbose: print(f'file [{index}]: {file_name_txt} was already processed')

        # if to_doc hasn't been done already
        if to_doc and (not os.path.exists(path_doc)):
            doc = docx.Document(FILE_TRANSCRIPTION_TEMPLATE)
            doc.add_heading(file_name_without_ext, 1)
            doc.add_paragraph(transcription_text)
            doc.save(path_doc)
            df.loc[index, 'transcription_doc'] = file_name_doc
            processed = True
        elif to_doc and os.path.exists(path_doc):
            if verbose: print(f'file [{index}]: {file_name_doc} was already processed')
                
        # add counter if processed 
        if processed: 
            if verbose: print(f'processed [{index}]: {file_name}')
            counter+=1

    # finish progress bar
    progressBar(1, 1, prefix=f'processing {len(df)} transcriptions', suffix='done.')
    
    # print message that we are done
    total_txt = len(df[df["transcription_txt"].notnull()])
    total_doc = len(df[df["transcription_doc"].notnull()])
    print(f'{counter} new transcriptions were processed, there are {total_txt} txt and {total_doc} doc transcriptions available')


def get_verified_transcriptions(df, path_transcriptions_json, path_transcriptions_txt, path_transcriptions_doc) -> pd.DataFrame:
    # find records to clean, we are probably waiting for a number of 
    # transcription responses - but the quota on webhooks.site might have been exceeded

    # do not make changes on the original
    df = df.copy()

    # loop over records and check them
    counter=0
    for index, row in df.iterrows():
        progressBar(index, len(df), prefix=f'verifying {len(df)} transcriptions')
        file_name_txt = row['transcription_txt']
        file_name_doc = row['transcription_doc']
        file_name_json = row['transcription_json']
        path_json = f'{path_transcriptions_json}/{file_name_json}'
        path_txt = f'{path_transcriptions_txt}/{file_name_txt}'
        path_doc = f'{path_transcriptions_doc}/{file_name_doc}'

        # initial condition
        valid = True

        # if json wasn't even retrieved, clean transcription_id,
        # the transcription will need to be re-obtained
        if pd.isna(file_name_json):
            valid = False
            df.loc[index, 'transcription_id'] = None
        
        # check jsons, clean json if file not present
        if not pd.isna(file_name_json):
            if not os.path.exists(path_json):
                valid = False
                df.loc[index, 'transcription_json'] = None

        if not pd.isna(file_name_txt):
            if not os.path.exists(path_txt):
                valid = False
                df.loc[index, 'transcription_txt'] = None

        if not pd.isna(file_name_doc):
            if not os.path.exists(path_doc):
                valid = False
                df.loc[index, 'transcription_doc'] = None

        # increment counter if all good
        if valid: counter+=1

    # finish progress bar
    progressBar(1, 1, prefix=f'verifying {len(df)} transcriptions', suffix='done.')
    
    # drop transcription_id column, it isn't needed in the final file
    df.drop('transcription_id', axis=1, inplace=True)
    print(f'there are {counter} processed and valid transcriptions')

    # if all transcriptions are done, print green message
    if len(df) == counter:
        coloured_print(f"*** all {len(df)} recordings were transcribed ***", colour='lightgreen')

    # return cleaned dataframe
    return df

# EOF
