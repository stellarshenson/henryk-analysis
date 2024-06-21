import requests
import pandas as pd
import os
import shutil
import io
import json
import math
import docx
import progress
import pathlib

import subprocess
from glob import glob


# local imports
from lib_henryk.logger import *
from lib_henryk.config import *
from lib_henryk import utils
from lib_henryk.recordings import recordings

def init(transcriptions_parquet_path:str = FILE_TRANSCRIPTIONS_PARQUET, path_recordings=DIR_RECORDINGS) -> (pd.DataFrame, pd.DataFrame):    
    # init transcriptions
    df_transcriptions = get_df_transcriptions(transcriptions_parquet_path)
    # generate transcriptions log
    df_transcriptions_log = get_df_transcriptions_log(df_transcriptions=df_transcriptions, path_recordings=path_recordings)
    return df_transcriptions_log, df_transcriptions 

def get_df_transcriptions_log(df_transcriptions: pd.DataFrame = None, path_recordings: str = DIR_RECORDINGS) -> pd.DataFrame:
    """create new empty dataframe or create a dataframe with information 
    abot existing transcriptions
    """
    # get list of files from the recordings directory
    files_list = recordings.get_recordings_files_list(path_recordings) # get list of files
    files_list = [os.path.basename(f) for f in files_list] # strip path
    names_list = ["".join(f.split(".")[:-1]) for f in files_list] # extract name

    # create initial transcriptions log
    df_transcriptions_log = pd.DataFrame( {
            'name': names_list, 
            'transcription_exists': [False] * len(names_list),
            'transcription_id': [None] * len(names_list),
    } )

    # update transcriptions log to show which transcriptions exist
    if df_transcriptions is not None:
        df_transcriptions = df_transcriptions[['name']].copy() # take only the files
        df_transcriptions.set_index('name', drop=True, inplace=True)  # set the file to be the update key (index)
        df_transcriptions['transcription_exists'] = True # broadcast that transcriptions exist

        df_transcriptions_log.set_index('name', drop=True, inplace=True) # set index to be the same as for transcriptions
        df_transcriptions_log.update(df_transcriptions, join='left', overwrite=True)

    # make sure that files are ordered in their natural temporal order (from name)
    df_transcriptions_log = recordings.sort_df_by_date_inferred_from_name(df_transcriptions_log)
    
    # reset index back to original and return
    return df_transcriptions_log

def get_df_transcriptions(transcriptions_parquet_path=FILE_TRANSCRIPTIONS_PARQUET) -> pd.DataFrame:
    # try to retrieve df_transcriptions if exist
    df_transcriptions = None
    if transcriptions_parquet_path != None and pathlib.Path(transcriptions_parquet_path).is_file():
        df_transcriptions = pd.read_parquet(transcriptions_parquet_path)

        # make sure transcriptions are ordered by date
        df_transcriptions = recordings.sort_df_by_date_inferred_from_name(df_transcriptions)
    else:
        df_transcriptions = pd.DataFrame( {
        'name': [], 
        'transcription': [],
    } )
    return df_transcriptions

def submit_transcriptions_goodtape( 
    webhooks_token_id, 
    goodtape_api_key, 
    df_transcriptions_log: pd.DataFrame = None, 
    df_transcriptions: pd.DataFrame = None, 
    path_recordings=DIR_RECORDINGS, 
    start=0, 
    stop=None, 
    verbose=False,
):  
    # make a copy or create new transcriptions log
    if df_transcriptions_log is not None:
        df_transcriptions_log = df_transcriptions_log.copy()
    else:
        df_transcriptions_log = get_df_transcriptions_log(df_transcriptions)
    
    # set up callback
    token_id = webhooks_token_id
    callback_url = f'https://webhook.site/{token_id}'

    # correct stop and start
    start = start if (start is not None) and (start >= 0) and (start < len(df_transcriptions_log)) else 0
    stop = stop if (stop is not None) and (stop > 0) and (stop <= len(df_transcriptions_log)) and (stop <= start) else len(df_transcriptions_log)

    # generate file name df for name resolution
    recordings_df = recordings.get_recordings_names_df(path_recordings).set_index('name', drop=True)
    
    # load just few records
    counter=0
    for index, row in df_transcriptions_log[start:stop].iterrows():
        # run transcription API if needed, only if transcription wasn't done already
        # we are only checking for json file, because transcription_id won't be needed anymore
        if row['transcription_exists'] == False:
            name = row['name']
            
            # this is to show which file is submitted
            trunc_name = utils.get_truncated_string(name, max_length=100)
            print(f'processing [{index}]: {trunc_name}')

            # resolve file name
            file_name = recordings_df.loc[name]['file']
            file_path = recordings_df.loc[name]['path']
            
            # copy file to /tmp, curl has issue with some paths
            temp_ext = file_name.split('.')[-1]
            temp_path = f'/tmp/recording.{temp_ext}'
            shutil.copy(file_path, temp_path)

            # perform transcription via API & CURL        
            response = subprocess.check_output(
                f'curl -s -X POST "https://api.goodtape.io/transcribe"' \
                + f' -H "Authorization: {goodtape_api_key}"' \
                + f' -F "audio=@{temp_path}"' \
                + f' -F "callbackUrl={callback_url}"' \
                + f' -F "languageCode=pl"' \
                + f' -F "speakerLabels=true"' \
                + f' -F "timeStamps=false"' 
            , shell=True, ) 

            # verify and write down the response
            try:
                response_json = json.loads(response)
                transcription_id = response_json['transcriptionId']
                df_transcriptions_log.loc[index, 'transcription_id'] = transcription_id
                if verbose:
                    print(f'file [{index}]: {transcription_id}')
                counter+=1
            except:
                print(response)
                break
        else:
            if verbose:
                print(f'file [{index}]: {file_name} already submitted')
    
    print(f'submitted {counter} recordings to the transcription service GoodTape')
    return df_transcriptions_log


def retrieve_responses_goodtape_via_webhooks(
    df_transcriptions_log: pd.DataFrame,
    df_transcriptions: pd.DataFrame,
    webhooks_token_id, 
) -> (pd.DataFrame, pd.DataFrame):
    # sanity check
    if df_transcriptions_log is None or df_transcriptions is None:
        raise Exception("transcriptions log doesn't exist, call `transcriptions.init()` to get them")
    
    # copy dataframe
    df_transcriptions_log = df_transcriptions_log.copy()
    df_transcriptions = df_transcriptions.copy()
    
    # use webhooks.site API to pull latest requests list
    headers = {}
    r = requests.get('https://webhook.site/token/'+ webhooks_token_id +'/requests?sorting=newest', headers=headers)
    print(f'found {len(r.json()["data"])} requests to fetch')
    
    # process requests one by one
    counter=0
    for i, request in enumerate(r.json()['data']):
        # fetch the data         
        response = requests.get(f'https://webhook.site/token/{webhooks_token_id}/request/{request["uuid"]}/raw')
        transcription_json = response.json()
        transcription_id = transcription_json['transcription_id']

        # before we write the file, let's just check if
        # we have a corresponding transcription_id in the database
        idx = df_transcriptions_log[df_transcriptions_log['transcription_id']==transcription_id].index
        if idx == None: 
            continue
        else:
            name = df_transcriptions_log.loc[idx[0]]['name']
            name_index = idx[0]
            name_truncated = utils.get_truncated_string(name, max_length=64)

            # write to transcriptions dataframe
            transcription = transcription_json['content']['text'] # retrieve transcription from json
            transcription = transcription.strip() # strip surrounding whitespaces
            row = {'name': name,'transcription': transcription} # generate df row
            df_transcriptions.loc[len(df_transcriptions)] = row # insert row
    
        # mark as done in the transcription log
        df_transcriptions_log.loc[name_index, 'transcription_exists'] = True
        print(f'retrieved transcription [{i}]: {transcription_id} -> file [{name_index}]: {name_truncated}')
        counter+=1
        
        # clean up this request as it was fetched
        response = requests.delete(f'https://webhook.site/token/{webhooks_token_id}/request/{request["uuid"]}')
    
    # print message that we are done
    print(f'{counter} transcriptions were retrieved')

    # check if we have anything more to wait for
    if df_transcriptions_log[(df_transcriptions_log['transcription_id'] == None) & (df_transcriptions_log['transcription_id'] == '')].empty:
        coloured_print(f"*** all submitted recordings were transcribed ***", colour='lightgreen')

    # return both dataframes - log and transcriptions
    return df_transcriptions_log, df_transcriptions


def transcriptions_to_doc(
    df_transcriptions: pd.DataFrame, 
    path_transcriptions_doc: str, 
    doc_template_path: str = FILE_TRANSCRIPTION_TEMPLATE,
    overwrite = True,
    verbose=False
):
    # iterate over records and process files
    counter=0
    for index, row in df_transcriptions.iterrows():
        progressBar(index, len(df_transcriptions), prefix=f'processing {len(df_transcriptions)} transcriptions')
        
        name = row['name']

        # doc and file names and paths
        path_doc = f'{path_transcriptions_doc}/{name}.docx'

        # open json file
        transcription = row['transcription']

        # if doc file is a new one - increment a counter
        if not pathlib.Path(path_doc).is_file():
            counter+=1
        
        # if doc file exists - it will be overwritten if set
        if not pathlib.Path(path_doc).is_file() or overwrite:
            doc = docx.Document(doc_template_path)
            doc.add_heading(name, 1)
            doc.add_paragraph(transcription)
            doc.save(path_doc)

    # finish progress bar
    progressBar(1, 1, prefix=f'processing {len(df_transcriptions)} transcriptions', suffix='done.')
    
    # print message that we are done, assume all transcriptions were processed
    total_doc = len(df_transcriptions)
    print(f'{counter} new transcriptions were processed, there are {total_doc} doc transcriptions available')


def get_verified_transcriptions(df_transcriptions_log, df_transcriptions) -> pd.DataFrame:
    # find records to clean, we are probably waiting for a number of 
    # transcription responses - but the quota on webhooks.site might have been exceeded

    # do not make changes on the original
    df_transcriptions_log = df_transcriptions_log.copy()
    
    # loop over records and check if thranscriptions were generated
    counter=0
    for index, row in df_transcriptions_log.iterrows():
        progressBar(index, len(df_transcriptions_log), prefix=f'verifying {len(df_transcriptions_log)} transcriptions')

        # retrieve name as key
        name = row['name']

        # initial condition
        valid = True

        # if json wasn't even retrieved, clean transcription_id,
        # the transcription will need to be re-obtained
        if df_transcriptions[df_transcriptions['name'] == name].empty:
            valid = False
            df.loc[index, 'transcription_id'] = None

        # increment counter if all good
        if valid: counter+=1

    # finish progress bar
    progressBar(1, 1, prefix=f'verifying {len(df_transcriptions_log)} transcriptions', suffix='done.')
    
    # clean transcription_id column, it needs to be empty if we need to resubmit
    df_transcriptions_log['transcription_id'] = None
    print(f'there are {counter} out of {len(df_transcriptions_log)} valid transcriptions')

    # if all transcriptions are done, print green message
    if len(df_transcriptions_log) == counter:
        coloured_print(f"*** all {len(df_transcriptions_log)} recordings were transcribed ***", colour='lightgreen')

    # return cleaned dataframe
    return df_transcriptions_log

def generate_transciptions_df_from_json_files(path_transcriptions_json: str) -> pd.DataFrame:
    files_json = glob(f'{path_transcriptions_json}/**/*.json', recursive=True)
    # files_json = [os.path.basename(f) for f in files_json]

    # iterate over files and produce entries for the dataset
    df_transcriptions = pd.DataFrame( {
        'name': [], 
        'transcription': [],
    } )

    # iterate over jsons and add to df
    for i, f in enumerate(files_json):
        progressBar(i,len(files_json)-1, prefix=f'processing {len(files_json)} json files')

        # fill in json, txt and doc if found
        file_name = os.path.basename(f) 
        name = "".join(file_name.split(".")[:-1])

        # read json
        with open(f, 'r') as f:
            transcription_json = json.load(f)
            f.close()

        transcription = transcription_json['content']['text']
        transcription = transcription.strip()

        # create a row to add to dataframe
        row = {
            'name': name, # file name without path, used as key 
            'transcription': transcription, # just the transcription text 
        }
    
        # add row to dataframe
        if row != None:
             df_transcriptions.loc[len(df_transcriptions)] = row
    
    # return dataframe
    return df_transcriptions

# EOF