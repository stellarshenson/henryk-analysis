"""
Transcription processing for recordings.
"""

from glob import glob
import json
import os
import pathlib
import shutil
import subprocess

import docx
import pandas as pd
import requests

from henryk_analysis import utils
from henryk_analysis.config import (
    DIR_RECORDINGS,
    FILE_TRANSCRIPTION_TEMPLATE,
    FILE_TRANSCRIPTIONS_PARQUET,
)
from henryk_analysis.logger import coloured_print, progressBar
from henryk_analysis.recordings import recordings


def init(
    transcriptions_parquet_path: str | pathlib.Path = FILE_TRANSCRIPTIONS_PARQUET,
    path_recordings: str | pathlib.Path = DIR_RECORDINGS,
) -> tuple[pd.DataFrame, pd.DataFrame]:
    """Initialize transcriptions and transcriptions log DataFrames."""
    df_transcriptions = get_df_transcriptions(transcriptions_parquet_path)
    df_transcriptions_log = get_df_transcriptions_log(
        df_transcriptions=df_transcriptions, path_recordings=path_recordings
    )
    return df_transcriptions_log, df_transcriptions


def get_df_transcriptions_log(
    df_transcriptions: pd.DataFrame = None,
    path_recordings: str | pathlib.Path = DIR_RECORDINGS,
) -> pd.DataFrame:
    """Create DataFrame with information about existing transcriptions."""
    files_list = recordings.get_recordings_files_list(path_recordings)
    files_list = [os.path.basename(f) for f in files_list]
    names_list = [".".join(f.split(".")[:-1]) for f in files_list]

    df_transcriptions_log = pd.DataFrame(
        {
            "name": names_list,
            "transcription_exists": [False] * len(names_list),
            "transcription_id": [None] * len(names_list),
        }
    )

    if df_transcriptions is not None:
        df_transcriptions = df_transcriptions[["name"]].copy()
        df_transcriptions.set_index("name", drop=True, inplace=True)
        df_transcriptions["transcription_exists"] = True

        df_transcriptions_log.set_index("name", drop=True, inplace=True)
        df_transcriptions_log.update(df_transcriptions, join="left", overwrite=True)

    df_transcriptions_log = recordings.sort_df_by_date_inferred_from_name(df_transcriptions_log)
    return df_transcriptions_log


def get_df_transcriptions(
    transcriptions_parquet_path: str | pathlib.Path = FILE_TRANSCRIPTIONS_PARQUET,
) -> pd.DataFrame:
    """Retrieve transcriptions DataFrame from parquet file."""
    transcriptions_parquet_path = pathlib.Path(transcriptions_parquet_path)
    if transcriptions_parquet_path.is_file():
        df_transcriptions = pd.read_parquet(transcriptions_parquet_path)
        df_transcriptions = recordings.sort_df_by_date_inferred_from_name(df_transcriptions)
    else:
        df_transcriptions = pd.DataFrame(
            {
                "name": [],
                "transcription": [],
            }
        )
    return df_transcriptions


def submit_transcriptions_goodtape(
    webhooks_token_id: str,
    goodtape_api_key: str,
    df_transcriptions_log: pd.DataFrame = None,
    df_transcriptions: pd.DataFrame = None,
    path_recordings: str | pathlib.Path = DIR_RECORDINGS,
    start: int = 0,
    stop: int = None,
    verbose: bool = False,
) -> pd.DataFrame:
    """Submit recordings for transcription via GoodTape API."""
    if df_transcriptions_log is not None:
        df_transcriptions_log = df_transcriptions_log.copy()
    else:
        df_transcriptions_log = get_df_transcriptions_log(df_transcriptions)

    token_id = webhooks_token_id
    callback_url = f"https://webhook.site/{token_id}"

    start = (
        start
        if (start is not None) and (start >= 0) and (start < len(df_transcriptions_log))
        else 0
    )
    stop = (
        stop
        if (stop is not None)
        and (stop > 0)
        and (stop <= len(df_transcriptions_log))
        and (stop <= start)
        else len(df_transcriptions_log)
    )

    recordings_df = recordings.get_recordings_names_df(path_recordings).set_index(
        "name", drop=True
    )

    counter = 0
    for index, row in df_transcriptions_log[start:stop].iterrows():
        if row["transcription_exists"] is False:
            name = row["name"]

            trunc_name = utils.get_truncated_string(name, max_length=100)
            print(f"processing [{index}]: {trunc_name}")

            file_name = recordings_df.loc[name]["file"]
            file_path = recordings_df.loc[name]["path"]

            temp_ext = file_name.split(".")[-1]
            temp_path = f"/tmp/recording.{temp_ext}"
            shutil.copy(file_path, temp_path)

            response = subprocess.check_output(
                f'curl -s -X POST "https://api.goodtape.io/transcribe"'
                f' -H "Authorization: {goodtape_api_key}"'
                f' -F "audio=@{temp_path}"'
                f' -F "callbackUrl={callback_url}"'
                f' -F "languageCode=pl"'
                f' -F "speakerLabels=true"'
                f' -F "timeStamps=false"',
                shell=True,
            )

            try:
                response_json = json.loads(response)
                transcription_id = response_json["transcriptionId"]
                df_transcriptions_log.loc[index, "transcription_id"] = transcription_id
                if verbose:
                    print(f"file [{index}]: {transcription_id}")
                counter += 1
            except Exception:
                print(response)
                break
        else:
            if verbose:
                print(f"file [{index}]: {file_name} already submitted")

    print(f"submitted {counter} recordings to the transcription service GoodTape")
    return df_transcriptions_log


def retrieve_responses_goodtape_via_webhooks(
    df_transcriptions_log: pd.DataFrame,
    df_transcriptions: pd.DataFrame,
    webhooks_token_id: str,
) -> tuple[pd.DataFrame, pd.DataFrame]:
    """Retrieve transcription responses from webhooks.site."""
    if df_transcriptions_log is None or df_transcriptions is None:
        raise ValueError(
            "transcriptions log doesn't exist, call `transcriptions.init()` to get them"
        )

    df_transcriptions_log = df_transcriptions_log.copy()
    df_transcriptions = df_transcriptions.copy()

    r = requests.get(
        f"https://webhook.site/token/{webhooks_token_id}/requests?sorting=newest",
        headers={},
    )
    print(f"found {len(r.json()['data'])} requests to fetch")

    counter = 0
    for i, request in enumerate(r.json()["data"]):
        response = requests.get(
            f"https://webhook.site/token/{webhooks_token_id}/request/{request['uuid']}/raw"
        )
        transcription_json = response.json()
        transcription_id = transcription_json["transcription_id"]

        idx = df_transcriptions_log[
            df_transcriptions_log["transcription_id"] == transcription_id
        ].index
        if idx is None:
            continue
        else:
            name = df_transcriptions_log.loc[idx[0]]["name"]
            name_index = idx[0]
            name_truncated = utils.get_truncated_string(name, max_length=64)

            transcription = transcription_json["content"]["text"]
            transcription = transcription.strip()
            row = {"name": name, "transcription": transcription}
            df_transcriptions.loc[len(df_transcriptions)] = row

        df_transcriptions_log.loc[name_index, "transcription_exists"] = True
        print(
            f"retrieved transcription [{i}]: {transcription_id} -> file [{name_index}]: {name_truncated}"
        )
        counter += 1

        requests.delete(
            f"https://webhook.site/token/{webhooks_token_id}/request/{request['uuid']}"
        )

    print(f"{counter} transcriptions were retrieved")

    if df_transcriptions_log[
        (df_transcriptions_log["transcription_id"] == None)  # noqa: E711
        & (df_transcriptions_log["transcription_id"] == "")
    ].empty:
        coloured_print("*** all submitted recordings were transcribed ***", colour="lightgreen")

    return df_transcriptions_log, df_transcriptions


def transcriptions_to_doc(
    df_transcriptions: pd.DataFrame,
    path_transcriptions_doc: str | pathlib.Path,
    doc_template_path: str | pathlib.Path = FILE_TRANSCRIPTION_TEMPLATE,
    overwrite: bool = True,
    verbose: bool = False,
) -> None:
    """Convert transcriptions DataFrame to Word documents."""
    counter = 0
    for index, row in df_transcriptions.iterrows():
        progressBar(
            index,
            len(df_transcriptions),
            prefix=f"processing {len(df_transcriptions)} transcriptions",
        )

        name = row["name"]
        path_doc = f"{path_transcriptions_doc}/{name}.docx"

        transcription = row["transcription"]

        if not pathlib.Path(path_doc).is_file():
            counter += 1

        if not pathlib.Path(path_doc).is_file() or overwrite:
            doc = docx.Document(str(doc_template_path))
            doc.add_heading(name, 1)
            doc.add_paragraph(transcription)
            doc.save(path_doc)

    progressBar(
        1,
        1,
        prefix=f"processing {len(df_transcriptions)} transcriptions",
        suffix="done.",
    )

    total_doc = len(df_transcriptions)
    print(
        f"{counter} new transcriptions were processed, there are {total_doc} doc transcriptions available"
    )


def get_verified_transcriptions(
    df_transcriptions_log: pd.DataFrame, df_transcriptions: pd.DataFrame
) -> pd.DataFrame:
    """Verify transcriptions and clean up missing ones."""
    df_transcriptions_log = df_transcriptions_log.copy()

    counter = 0
    for index, row in df_transcriptions_log.iterrows():
        progressBar(
            index,
            len(df_transcriptions_log),
            prefix=f"verifying {len(df_transcriptions_log)} transcriptions",
        )

        name = row["name"]
        valid = True

        if df_transcriptions[df_transcriptions["name"] == name].empty:
            valid = False
            df_transcriptions_log.loc[index, "transcription_id"] = None

        if valid:
            counter += 1

    progressBar(
        1,
        1,
        prefix=f"verifying {len(df_transcriptions_log)} transcriptions",
        suffix="done.",
    )

    df_transcriptions_log["transcription_id"] = None
    print(f"there are {counter} out of {len(df_transcriptions_log)} valid transcriptions")

    if len(df_transcriptions_log) == counter:
        coloured_print(
            f"*** all {len(df_transcriptions_log)} recordings were transcribed ***",
            colour="lightgreen",
        )

    return df_transcriptions_log


def generate_transcriptions_df_from_json_files(
    path_transcriptions_json: str | pathlib.Path,
) -> pd.DataFrame:
    """Generate transcriptions DataFrame from JSON files."""
    path_transcriptions_json = str(path_transcriptions_json)
    files_json = glob(f"{path_transcriptions_json}/**/*.json", recursive=True)

    df_transcriptions = pd.DataFrame(
        {
            "name": [],
            "transcription": [],
        }
    )

    for i, f in enumerate(files_json):
        progressBar(i, len(files_json) - 1, prefix=f"processing {len(files_json)} json files")

        file_name = os.path.basename(f)
        name = ".".join(file_name.split(".")[:-1])

        with open(f, "r") as json_file:
            transcription_json = json.load(json_file)

        transcription = transcription_json["content"]["text"]
        transcription = transcription.strip()

        row = {
            "name": name,
            "transcription": transcription,
        }

        if row is not None:
            df_transcriptions.loc[len(df_transcriptions)] = row

    return df_transcriptions


# EOF
